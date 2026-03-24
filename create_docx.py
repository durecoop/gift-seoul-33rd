from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)


def styled_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return h


def cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def cell_text(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_run(p, text, bold=False, size=11, color=None, italic=False):
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for r in p.runs:
        r.font.name = '맑은 고딕'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')


WHITE = (0xFF, 0xFF, 0xFF)
GRAY = (0x66, 0x66, 0x66)
LIGHT_GRAY = (0x99, 0x99, 0x99)

# ═══════════════════════════════════
# 표지
# ═══════════════════════════════════
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '서울남부두레생협', bold=True, size=14, color=(0x2d, 0x6a, 0x4f))

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '퇴계점 폐점 후\n온라인쇼핑몰 전환 증정 기획전', bold=True, size=26)

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '"매장은 문을 닫아도, 조합원님과의 인연은 계속됩니다"', size=13, italic=True, color=(0x88, 0x88, 0x88))

for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '2026년 3월 24일\n기획팀 / 온라인사업팀', size=11, color=GRAY)

doc.add_page_break()

# ═══════════════════════════════════
# 1. 기획 배경
# ═══════════════════════════════════
styled_heading('1. 기획 배경', level=1)
styled_heading('1-1. 현황', level=2)

p = doc.add_paragraph()
add_run(p, '퇴계점 폐점일: ', bold=True)
add_run(p, '2026년 3월 25일(수)')

bullet('퇴계점 이용 조합원의 오프라인 구매 채널 소멸')
bullet('대체 매장까지의 거리 및 접근성 문제로 이탈 우려')
bullet('온라인쇼핑몰 미이용 조합원 비율이 높은 점포 (50~60대 조합원 다수)')

styled_heading('1-2. 기획 목적', level=2)

for title, desc in [
    ('구매 이탈 방지', '폐점 후 구매 공백기를 최소화하고, 조합원의 생활재 구매를 지속 가능하게 합니다.'),
    ('온라인 전환 유도', '첫 온라인 주문 경험을 긍정적으로 만들어 자연스러운 채널 전환을 이끌어냅니다.'),
    ('정기 주문 습관 형성', '5주간 매주 1가지씩 순차 증정하여 온라인 주문이 일상이 되도록 합니다.'),
    ('혜택 형평성 유지', '서울남부 33주년 기획전과 동일한 품목을 증정하여 매장 간 형평성을 지킵니다.'),
]:
    p = doc.add_paragraph()
    add_run(p, f'  {title}: ', bold=True)
    add_run(p, desc)

styled_heading('1-3. 생협으로서의 의미', level=2)

p = doc.add_paragraph()
add_run(p, (
    '두레생협은 33년간 "생산자와 소비자가 서로 얼굴을 아는 관계"를 지켜왔습니다. '
    '퇴계점 폐점은 물리적 공간의 변화이지, 이 관계의 단절이 아닙니다. '
    '온라인쇼핑몰은 매장을 대체하는 것이 아니라, 조합원과 생산자를 잇는 또 하나의 다리입니다.\n\n'
    '이번 증정 기획전은 단순한 판촉이 아닌, '
    '"함께 나누는 생협의 가치"를 온라인에서도 이어가겠다는 약속입니다. '
    '특히 50~60대 조합원님들이 온라인에 대한 부담 없이 편하게 이용하실 수 있도록, '
    '주문이 어려우신 분께는 두레생협 상담실(1661-5110)을 통한 지원도 함께 안내합니다.'
))

styled_heading('1-4. 타겟 조합원 특성', level=2)

bullet('50~60대 여성 조합원이 주 이용층')
bullet('온라인 주문 경험이 적거나 처음이신 분 다수')
bullet('매장에서의 대면 구매에 익숙하시며, 신뢰 관계를 중시')
bullet('가족(자녀) 도움으로 주문 가능 - "자녀분이 대신 주문해 주셔도 됩니다" 안내 필요')
bullet('전화 상담을 통한 주문 지원 필수 (두레생협 상담실 1661-5110)')

doc.add_page_break()

# ═══════════════════════════════════
# 2. 증정 개요
# ═══════════════════════════════════
styled_heading('2. 증정 개요', level=1)

table = doc.add_table(rows=9, cols=2, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

overview = [
    ('구분', '내용'),
    ('이벤트명', '퇴계점 조합원님께 드리는 온라인 특별 증정'),
    ('기간', '2026. 3. 30(월) ~ 5. 1(금) [5주간 / 14주차~18주차]'),
    ('대상', '퇴계점 이용 조합원 (온라인쇼핑몰 주문 시)'),
    ('증정 방식', '매주 1품목씩 순차 증정 (5주간 5품목)'),
    ('최소 주문금액', '50,000원 이상 (배송일 기준)'),
    ('증정 한도', '1인 1회/주 (중복 수령 불가)'),
    ('적용 기준', '주문일이 아닌 배송일 기준'),
    ('문의', '두레생협 상담실 1661-5110'),
]

for i, (c1, c2) in enumerate(overview):
    if i == 0:
        cell_text(table.cell(i, 0), c1, bold=True, size=10, color=WHITE)
        cell_text(table.cell(i, 1), c2, bold=True, size=10, color=WHITE)
        cell_shading(table.cell(i, 0), '2D6A4F')
        cell_shading(table.cell(i, 1), '2D6A4F')
    else:
        cell_text(table.cell(i, 0), c1, bold=True, size=10)
        cell_text(table.cell(i, 1), c2, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        cell_shading(table.cell(i, 0), 'F0F7F0')

for row in table.rows:
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(12)

doc.add_paragraph('')

# ═══════════════════════════════════
# 3. 주차별 증정 품목
# ═══════════════════════════════════
styled_heading('3. 주차별 증정 품목 (서울남부 33주년 기획전 동일)', level=1)

p = doc.add_paragraph()
add_run(p, '서울남부에서 요일별로 1품목씩 증정하는 5가지 품목을, 퇴계점 조합원님께는 매주 1품목씩 5주간 순차 증정합니다. 모든 품목 동일하게 5만원 이상 주문 시 증정됩니다.')

doc.add_paragraph('')

items_t = doc.add_table(rows=7, cols=5, style='Table Grid')
items_t.alignment = WD_TABLE_ALIGNMENT.CENTER

items = [
    ('차수', '주차(기간)', '품목명', '규격', '특징'),
    ('1차', '14주 (3/30~4/3)', '느타리버섯', '200g', '무농약 인증 / 봄 제철 버섯'),
    ('2차', '15주 (4/6~4/10)', '당근', '500g', '무농약 인증 / 생산자 직접 재배'),
    ('3차', '16주 (4/13~4/17)', '쑥버무리', '250g', '제철 간식 / 봄 향기 가득'),
    ('4차', '17주 (4/20~4/24)', '원삼유정란', '10구', '건강한 유정란'),
    ('5차', '18주 (4/27~5/1)', '냉장 손질삼치', '300g', '손질 완료 / 신선 냉장'),
    ('', '', '공통 조건', '', '5만원 이상 주문 시 해당 주 품목 자동 동봉'),
]

for i, row_data in enumerate(items):
    for j, val in enumerate(row_data):
        if i == 0:
            cell_text(items_t.cell(i, j), val, bold=True, size=9, color=WHITE)
            cell_shading(items_t.cell(i, j), '5C3D2E')
        elif i == 6:
            cell_text(items_t.cell(i, j), val, bold=True, size=9)
            cell_shading(items_t.cell(i, j), 'FAF3EB')
        else:
            bold = (j == 2)
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 4 else WD_ALIGN_PARAGRAPH.CENTER
            cell_text(items_t.cell(i, j), val, bold=bold, size=9, align=align)

doc.add_paragraph('')

p = doc.add_paragraph()
add_run(p, '※ 순차 증정 방식의 장점: ', bold=True, size=10)
add_run(p, (
    '매주 새로운 선물이 있으므로 조합원님이 5주간 지속적으로 주문하게 되어, '
    '온라인 장보기 습관이 자연스럽게 형성됩니다.'
), size=10)

doc.add_page_break()

# ═══════════════════════════════════
# 4. 운영 일정
# ═══════════════════════════════════
styled_heading('4. 운영 일정', level=1)

sched_t = doc.add_table(rows=6, cols=3, style='Table Grid')
sched_t.alignment = WD_TABLE_ALIGNMENT.CENTER

sched = [
    ('차수', '기간', '증정 품목'),
    ('1차 (14주)', '3/30(월) ~ 4/3(금)', '느타리버섯 200g'),
    ('2차 (15주)', '4/6(월) ~ 4/10(금)', '당근 500g'),
    ('3차 (16주)', '4/13(월) ~ 4/17(금)', '쑥버무리 250g'),
    ('4차 (17주)', '4/20(월) ~ 4/24(금)', '원삼유정란 10구'),
    ('5차 (18주)', '4/27(월) ~ 5/1(금)', '냉장 손질삼치 300g (마지막)'),
]

for i, row_data in enumerate(sched):
    for j, val in enumerate(row_data):
        if i == 0:
            cell_text(sched_t.cell(i, j), val, bold=True, size=10, color=WHITE)
            cell_shading(sched_t.cell(i, j), '2D6A4F')
        else:
            bold = (j == 0)
            align = WD_ALIGN_PARAGRAPH.LEFT if j >= 1 else WD_ALIGN_PARAGRAPH.CENTER
            cell_text(sched_t.cell(i, j), val, bold=bold, size=10, align=align)

doc.add_paragraph('')

# ═══════════════════════════════════
# 5. 홍보 계획
# ═══════════════════════════════════
styled_heading('5. 홍보 계획', level=1)
styled_heading('5-1. 핵심 메시지 (50~60대 여성 조합원 대상)', level=2)

for title, msg in [
    ('메인 카피', '"매장은 문을 닫아도, 우리 가족 밥상을 지키는 마음은 변하지 않습니다"'),
    ('서브 카피', '"매주 새로운 선물이 집 앞까지 찾아갑니다 - 5주간 5가지 봄 선물"'),
    ('생협 가치', '"생산자와 소비자가 서로 얼굴을 아는 관계 - 온라인에서도 그대로입니다"'),
    ('안심 유도', '"온라인이 처음이셔도 걱정 마세요. 자녀분이 대신 주문해 주셔도 됩니다"'),
    ('상담 안내', '"주문이 어려우시면 두레생협 상담실 1661-5110으로 전화주세요"'),
]:
    p = doc.add_paragraph()
    add_run(p, f'  {title}\n', bold=True)
    add_run(p, f'   {msg}')

doc.add_paragraph('')
styled_heading('5-2. 홍보 강조 포인트', level=2)

for pt in [
    '매주 1가지씩 새로운 선물! 5주간 기다리는 즐거움',
    '무거운 장바구니 없이, 집 앞까지 편하게 배송',
    '매장에서 고르시던 바로 그 생활재, 품질 그대로',
    '된장찌개, 나물, 계란말이... 우리 가족 밥상에 필요한 것만 골랐어요',
    '5만원 이상이면 OK! 부담 없는 조건',
    '자녀분이 대신 주문해 주셔도 선물은 동일하게 동봉',
    '주문이 어려우시면 1661-5110 전화 한 통이면 됩니다',
]:
    bullet(pt)

doc.add_paragraph('')
styled_heading('5-3. 채널별 홍보 계획', level=2)

promo_t = doc.add_table(rows=7, cols=4, style='Table Grid')
promo_t.alignment = WD_TABLE_ALIGNMENT.CENTER

promo = [
    ('채널', '내용', '시기', '담당'),
    ('개별 문자/알림톡', '퇴계점 조합원 대상 폐점 안내 +\n온라인 증정 혜택 + 상담실 번호', '3/25 전후', '홍보팀'),
    ('매장 안내문', '폐점 전 매장 내 포스터/리플릿\n큰 글씨 + 쇼핑몰 QR코드 + 상담실 번호', '~3/25', '매장운영팀'),
    ('온라인 배너', '쇼핑몰 메인 페이지 배너\n퇴계점 조합원 타겟 팝업', '3/30~', '온라인사업팀'),
    ('카드뉴스(SNS)', '주차별 카드뉴스 5종 + 통합 카드뉴스\n생협 가치 + 따뜻한 어조', '매주 게시', '홍보팀'),
    ('조합원 소식지', '4월호 소식지 특집 게재', '4월호', '편집팀'),
    ('전화 안내', '온라인 미가입 조합원 개별 전화\n가입/주문 방법 안내 + 자녀 안내 권유', '3/25~3/31', '매장운영팀'),
]

for i, row_data in enumerate(promo):
    for j, val in enumerate(row_data):
        if i == 0:
            cell_text(promo_t.cell(i, j), val, bold=True, size=9, color=WHITE)
            cell_shading(promo_t.cell(i, j), '5C3D2E')
        else:
            bold = (j == 0)
            align = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 2, 3] else WD_ALIGN_PARAGRAPH.LEFT
            cell_text(promo_t.cell(i, j), val, bold=bold, size=9, align=align)

doc.add_page_break()

# ═══════════════════════════════════
# 6. 홍보 문안
# ═══════════════════════════════════
styled_heading('6. 홍보 문안', level=1)
styled_heading('6-1. 문자/알림톡 문안', level=2)

p = doc.add_paragraph()
add_run(p, (
    '[서울남부두레생협]\n\n'
    '퇴계점에서 장 보시던 조합원님,\n'
    '그동안 감사했습니다.\n\n'
    '3/25(수) 퇴계점 운영이 종료되지만,\n'
    '우리 가족 밥상을 지켜온\n'
    '두레의 마음은 변하지 않습니다.\n\n'
    '온라인쇼핑몰에서 5만원 이상 주문하시면\n'
    '매주 새로운 봄 선물 1가지를 보내드려요.\n\n'
    '1차 느타리버섯 / 2차 당근\n'
    '3차 쑥버무리 / 4차 유정란\n'
    '5차 손질삼치\n\n'
    '기간: 3/30~5/1 (5주간, 매주 1품목)\n'
    '쇼핑몰: ecoop.or.kr\n\n'
    '주문이 어려우시면\n'
    '상담실 1661-5110으로 전화주세요.\n'
    '자녀분이 대신 주문해 주셔도 됩니다.'
), size=10)

doc.add_paragraph('')
styled_heading('6-2. 매장 안내문 문안', level=2)

p = doc.add_paragraph()
add_run(p, (
    '[ 퇴계점 조합원님께 드리는 안내 ]\n\n'
    '사랑하는 조합원님께,\n\n'
    '퇴계점은 3월 25일(수)을 마지막으로\n'
    '운영을 종료하게 되었습니다.\n\n'
    '그동안 보내주신 사랑에 깊이 감사드립니다.\n\n'
    '매장 문은 닫히지만,\n'
    '매장에서 고르시던 바로 그 생활재를\n'
    '온라인에서도 만나실 수 있습니다.\n'
    '집 앞까지 편하게 배송해 드려요.\n\n'
    '온라인쇼핑몰에서 5만원 이상 주문하시면\n'
    '매주 1가지씩 선물을 보내드립니다.\n\n'
    '  1차(14주) 느타리버섯 200g\n'
    '  2차(15주) 당근 500g\n'
    '  3차(16주) 쑥버무리 250g\n'
    '  4차(17주) 원삼유정란 10구\n'
    '  5차(18주) 손질삼치 300g\n\n'
    '온라인 주문이 어려우시면\n'
    '두레생협 상담실로 전화주세요.\n'
    '친절하게 도와드리겠습니다.\n\n'
    '  상담실: 1661-5110\n'
    '  쇼핑몰: ecoop.or.kr\n\n'
    '33년간 지켜온 약속,\n'
    '앞으로도 조합원님의 밥상을 지키겠습니다.\n\n'
    '- 서울남부두레생협 임직원 일동'
), size=10)

doc.add_paragraph('')
styled_heading('6-3. SNS 카드뉴스 텍스트', level=2)

p = doc.add_paragraph()
add_run(p, (
    '[카드 1 - 커버]\n'
    '퇴계점 조합원님께 드리는 특별한 선물\n'
    '"매장은 문을 닫아도, 우리의 인연은 계속됩니다"\n\n'
    '[카드 2 - 편지]\n'
    '조합원님, 감사합니다.\n'
    '우리 가족 밥상을 지켜온 두레의 마음은 변하지 않습니다.\n'
    '매주 1가지씩, 5주간 선물을 준비했어요.\n\n'
    '[카드 3 - 이벤트 안내]\n'
    '매주 새로운 선물 1가지! 5주간 순차 증정\n'
    '조건: 5만원 이상 | 주문이 어려우시면 1661-5110\n\n'
    '[카드 4~8 - 주차별 품목]\n'
    '1차(14주) 느타리버섯 200g - 된장찌개 한 그릇 끓여보세요\n'
    '2차(15주) 당근 500g - 나물이든 주스든 어디에든\n'
    '3차(16주) 쑥버무리 250g - 손자 손녀와 함께 나눠 드세요\n'
    '4차(17주) 원삼유정란 10구 - 아침 식탁이 든든해져요\n'
    '5차(18주) 손질삼치 300g - 주말 밥상에 삼치구이 한 점\n\n'
    '[카드 9 - 마무리]\n'
    '우리 가족 밥상, 이제 온라인에서도 안심하세요\n'
    '주문이 어려우시면 1661-5110으로 전화주세요'
), size=10)

doc.add_page_break()

# ═══════════════════════════════════
# 7. 예산
# ═══════════════════════════════════
styled_heading('7. 예산 (5주 기준, 예상)', level=1)

p = doc.add_paragraph()
add_run(p, '주차별 1품목씩 순차 증정이므로, 각 주차에 해당 품목만 준비합니다.')

doc.add_paragraph('')

budget_t = doc.add_table(rows=8, cols=4, style='Table Grid')
budget_t.alignment = WD_TABLE_ALIGNMENT.CENTER

budget = [
    ('항목', '단가(추정)', '수량(해당 주)', '소계'),
    ('1차 느타리버섯 200g', '2,500원', '50개', '125,000원'),
    ('2차 당근 500g', '2,000원', '50개', '100,000원'),
    ('3차 쑥버무리 250g', '3,500원', '50개', '175,000원'),
    ('4차 원삼유정란 10구', '5,000원', '50개', '250,000원'),
    ('5차 냉장 손질삼치 300g', '6,000원', '50개', '300,000원'),
    ('홍보물 제작비', '-', '-', '200,000원'),
    ('총 예산', '', '', '1,150,000원'),
]

for i, row_data in enumerate(budget):
    for j, val in enumerate(row_data):
        if i == 0:
            cell_text(budget_t.cell(i, j), val, bold=True, size=9, color=WHITE)
            cell_shading(budget_t.cell(i, j), '2D6A4F')
        elif i == 7:
            cell_text(budget_t.cell(i, j), val, bold=True, size=10)
            cell_shading(budget_t.cell(i, j), 'D8F3DC')
        else:
            align = WD_ALIGN_PARAGRAPH.RIGHT if j >= 1 else WD_ALIGN_PARAGRAPH.LEFT
            cell_text(budget_t.cell(i, j), val, size=9, align=align)

doc.add_paragraph('')

p = doc.add_paragraph()
add_run(p, '※ 단가 및 수량은 예상치이며, 실제 원가 및 참여 규모에 따라 조정이 필요합니다.', size=9, color=LIGHT_GRAY)

doc.add_paragraph('')

# ═══════════════════════════════════
# 8. 기대 효과
# ═══════════════════════════════════
styled_heading('8. 기대 효과', level=1)

effect_t = doc.add_table(rows=5, cols=3, style='Table Grid')
effect_t.alignment = WD_TABLE_ALIGNMENT.CENTER

effects = [
    ('지표', '목표', '측정 방법'),
    ('퇴계점 조합원 온라인 전환율', '30% 이상', '조합원 DB 기반 온라인 첫 주문자 추적'),
    ('주간 증정 참여 건수', '주 50건 이상', '증정 실적 집계'),
    ('5주 내 재주문율', '60% 이상', '2회 이상 주문 조합원 비율'),
    ('평균 객단가', '5만원 이상 유지', '주문 금액 분석'),
]

for i, row_data in enumerate(effects):
    for j, val in enumerate(row_data):
        if i == 0:
            cell_text(effect_t.cell(i, j), val, bold=True, size=9, color=WHITE)
            cell_shading(effect_t.cell(i, j), '5C3D2E')
        else:
            bold = (j == 1)
            cell_text(effect_t.cell(i, j), val, bold=bold, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph('')

# ═══════════════════════════════════
# 9. 추진 일정
# ═══════════════════════════════════
styled_heading('9. 추진 일정', level=1)

timeline_t = doc.add_table(rows=9, cols=3, style='Table Grid')
timeline_t.alignment = WD_TABLE_ALIGNMENT.CENTER

timeline = [
    ('일정', '내용', '담당'),
    ('~3/25', '퇴계점 매장 내 안내물 게시 (큰 글씨 + 상담실 번호)', '매장운영팀'),
    ('3/25', '퇴계점 폐점 / 조합원 개별 문자 발송', '매장운영팀/홍보팀'),
    ('3/25~3/30', '온라인 미가입 조합원 전화 안내 + 가입 지원', '매장운영팀'),
    ('3/26~3/28', '온라인쇼핑몰 주차별 증정 세팅 및 테스트', '온라인사업팀'),
    ('3/28', '1차 카드뉴스 사전 홍보', '홍보팀'),
    ('3/30', '1차 증정 시작 (14주차 - 느타리버섯)', '온라인사업팀'),
    ('매주 월요일', '해당 주차 카드뉴스 게시 (2차~5차)', '홍보팀'),
    ('5/5~', '5주 성과 분석 및 후속 방안 결정', '기획팀'),
]

for i, row_data in enumerate(timeline):
    for j, val in enumerate(row_data):
        if i == 0:
            cell_text(timeline_t.cell(i, j), val, bold=True, size=9, color=WHITE)
            cell_shading(timeline_t.cell(i, j), '2D6A4F')
        else:
            bold = (j == 0)
            align = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 2] else WD_ALIGN_PARAGRAPH.LEFT
            cell_text(timeline_t.cell(i, j), val, bold=bold, size=9, align=align)

doc.add_paragraph('')

# ═══════════════════════════════════
# 10. 유의사항
# ═══════════════════════════════════
styled_heading('10. 운영 유의사항', level=1)

for title, desc in [
    ('1인 1회/주 제한', '동일 조합원 주당 1회만 증정하며, 중복 수령은 불가합니다.'),
    ('배송일 기준 적용', '주문일이 아닌 배송일 기준으로 해당 주차를 판단합니다.'),
    ('퇴계점 조합원 식별', '조합원 DB 기반으로 퇴계점 등록 조합원을 대상으로 자동 적용합니다.'),
    ('품목 소진 시', '해당 주차 증정을 조기 종료하며, 사전에 선착순 안내를 진행합니다.'),
    ('상담실 지원', '온라인 주문이 어려운 조합원을 위해 두레생협 상담실(1661-5110)을 통한 주문 지원을 안내합니다.'),
    ('가족 대리 주문', '자녀 등 가족이 대신 주문하더라도 조합원 본인 계정으로 주문 시 증정이 적용됩니다.'),
]:
    p = doc.add_paragraph()
    add_run(p, f'  {title}: ', bold=True)
    add_run(p, desc)

doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p, '2026년 3월 24일\n서울남부두레생협 기획팀 / 온라인사업팀', size=10, color=GRAY)

doc.save('퇴계점_온라인전환_증정기획전.docx')
print('Word 문서 생성 완료!')
